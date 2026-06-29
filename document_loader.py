import io
import re
import os
from PyPDF2 import PdfReader
from docx import Document
import streamlit as st

def load_and_split_document(uploaded_file, parent_size=1200, parent_overlap=150, child_size=300, child_overlap=50):
    """
    Load and split document into hierarchical chunks (parent-child structure).
    Supports PDF, DOCX, TXT, and MD formats.
    """
    try:
        file_type = uploaded_file.type
        file_name = uploaded_file.name.lower()
        
        # Extract text based on file type
        if file_type == "application/pdf" or file_name.endswith('.pdf'):
            text = extract_pdf_text(uploaded_file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith('.docx'):
            text = extract_docx_text(uploaded_file)
        elif file_type == "text/plain" or file_name.endswith('.txt') or file_name.endswith('.md') or file_type == "text/markdown":
            text = extract_txt_text(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        if not text.strip():
            raise ValueError("No text content found in the document")
        
        # Clean and preprocess text
        text = clean_text(text)
        
        # Split into parent chunks
        parent_chunks = split_text_with_overlap(text, parent_size, parent_overlap)
        
        # Build hierarchical chunks list
        hierarchical_chunks = []
        for parent_chunk in parent_chunks:
            page_number = find_page_number(text, parent_chunk)
            
            # Split parent chunk into child chunks
            child_chunks = split_text_with_overlap(parent_chunk, child_size, child_overlap)
            for child in child_chunks:
                if len(child.strip()) > 30:  # Filter out very short child chunks
                    hierarchical_chunks.append({
                        "child": child.strip(),
                        "parent": parent_chunk.strip(),
                        "page_number": page_number
                    })
        
        if not hierarchical_chunks:
            # Fallback if no hierarchical chunks created
            if len(text.strip()) > 0:
                hierarchical_chunks.append({
                    "child": text.strip()[:child_size],
                    "parent": text.strip(),
                    "page_number": 1
                })
            else:
                raise ValueError("No valid chunks created from the document")
        
        return hierarchical_chunks
        
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        raise e

def extract_pdf_text(uploaded_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
            except Exception as e:
                st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                continue
        
        return text
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")

def extract_docx_text(uploaded_file):
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(uploaded_file.read()))
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")

def extract_txt_text(uploaded_file):
    """Extract text from TXT file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                uploaded_file.seek(0)  # Reset file pointer
                text = uploaded_file.read().decode(encoding)
                return text
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    except Exception as e:
        raise ValueError(f"Error reading TXT: {str(e)}")

def clean_text(text):
    """Clean and preprocess text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep important punctuation
    text = re.sub(r'[^\w\s.,!?;:()\-"\']', ' ', text)
    
    # Fix common OCR errors
    text = re.sub(r'\b(\w+)\1+\b', r'\1', text)  # Remove repeated words
    
    # Normalize line breaks
    text = re.sub(r'\n+', '\n', text)
    
    # Remove page headers/footers (simple heuristic)
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if len(line) > 5 and not re.match(r'^(Page \d+|\d+)$', line):
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines).strip()

def split_text_with_overlap(text, chunk_size=800, overlap=100):
    """Split text into overlapping chunks"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Calculate end position
        end = start + chunk_size
        
        # If this is not the last chunk, try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings within a reasonable range
            sentence_break = find_sentence_break(text, end, chunk_size // 4)
            if sentence_break != -1:
                end = sentence_break
        
        # Extract chunk
        chunk = text[start:end].strip()
        
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        if end >= len(text):
            break
        
        start = max(start + chunk_size - overlap, start + 1)
    
    return chunks

def find_sentence_break(text, preferred_end, search_range):
    """Find the best sentence break near the preferred end position"""
    # Look for sentence endings
    sentence_endings = ['.', '!', '?', '\n']
    
    # Search backwards from preferred_end
    for i in range(min(search_range, preferred_end), 0, -1):
        pos = preferred_end - i
        if pos > 0 and text[pos] in sentence_endings:
            # Make sure it's not an abbreviation or decimal
            if text[pos] == '.' and pos > 0 and text[pos-1].isdigit():
                continue
            return pos + 1
    
    # Search forwards if no break found backwards
    for i in range(search_range):
        pos = preferred_end + i
        if pos < len(text) and text[pos] in sentence_endings:
            if text[pos] == '.' and pos > 0 and text[pos-1].isdigit():
                continue
            return pos + 1
    
    return -1

def get_document_preview(chunks, max_preview=500):
    """Get a preview of the document content"""
    if not chunks:
        return "No content available"
    
    preview = chunks[0]
    if len(preview) > max_preview:
        preview = preview[:max_preview] + "..."
    
    return preview

def validate_document(uploaded_file):
    """Validate uploaded document"""
    if uploaded_file is None:
        return False, "No file uploaded"
    
    # Check file size (10MB limit)
    max_size = 10 * 1024 * 1024  # 10MB
    if uploaded_file.size > max_size:
        return False, f"File too large. Maximum size: {max_size // (1024*1024)}MB"
    
    # Check file type/extension
    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "text/markdown"
    ]
    file_name = uploaded_file.name.lower()
    allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
    
    is_allowed = uploaded_file.type in allowed_types or any(file_name.endswith(ext) for ext in allowed_extensions)
    if not is_allowed:
        return False, f"Unsupported file type: {uploaded_file.type}"
    
    return True, "Valid document"

def find_page_number(full_text, chunk_text):
    """Find the page number that a chunk belongs to by searching backwards in full_text"""
    try:
        idx = full_text.find(chunk_text)
        if idx == -1:
            return 1
            
        preceding_text = full_text[:idx]
        matches = list(re.finditer(r'---\s*Page\s*(\d+)\s*---', preceding_text))
        if matches:
            return int(matches[-1].group(1))
            
        match_inside = re.search(r'---\s*Page\s*(\d+)\s*---', chunk_text)
        if match_inside:
            return int(match_inside.group(1))
            
        return 1
    except Exception:
        return 1

def save_pdf_locally(uploaded_file):
    """Save the uploaded PDF file to the data/uploaded_pdfs/ directory"""
    try:
        dir_path = "data/uploaded_pdfs"
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
            
        file_path = os.path.join(dir_path, uploaded_file.name)
        
        # Save bytes
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        with open(file_path, "wb") as f:
            f.write(file_bytes)
            
        # Reset file pointer for subsequent reads
        uploaded_file.seek(0)
        return True
    except Exception as e:
        st.error(f"Error saving PDF locally: {str(e)}")
        return False